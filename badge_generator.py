#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word Template Badge Generator for VIMC
Sá»­ dá»¥ng Word template cÃ³ sáºµn â†' Thay tháº¿ text (giá»¯ nguyÃªn format) â†' Export PNG
"""

import pandas as pd
from docx import Document
from docx2pdf import convert
from pdf2image import convert_from_path
import os
import shutil
import sys
from pathlib import Path
import platform
import re

# Ensure stdout/stderr use UTF-8 so Vietnamese accents print correctly on Windows consoles
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

REQUIRED_FONT_KEYWORDS = ['faustina']

# Common font directories to scan on Windows for installed fonts
WINDOWS_FONT_DIRS = [
    Path(os.environ.get('WINDIR', r'C:\Windows')) / 'Fonts',
    Path.home() / 'AppData' / 'Local' / 'Microsoft' / 'Windows' / 'Fonts',
]

# ===== Cáº¤U HÃŒNH =====
TEMPLATE_FILE = 'badge_template.docx'  # File Word template cá»§a báº¡n
EXCEL_FILE = 'danh_sach.xlsx'         # File Excel danh sÃ¡ch
OUTPUT_DIR = 'badges'                 # ThÆ° má»¥c output

# Táº¡o thÆ° má»¥c táº¡m vÃ  output
TEMP_DIR = 'temp_docs'
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)

def resolve_poppler_path():
    """
    Try to locate Poppler binaries.
    Checks POPPLER_PATH env, then bundled poppler folder inside project.
    """
    env_path = os.environ.get('POPPLER_PATH')
    if env_path:
        env_path = Path(env_path).expanduser()
        if env_path.exists():
            return str(env_path)
    local_poppler_root = Path(__file__).resolve().parent / 'poppler'
    if local_poppler_root.exists():
        for exe_name in ('pdftoppm.exe', 'pdftoppm'):
            matches = list(local_poppler_root.glob(f'**/{exe_name}'))
            if matches:
                return str(matches[0].parent)
    return None

POPPLER_PATH = resolve_poppler_path()

def find_missing_fonts():
    """
    Check for required fonts so rendered badges match the expected styling.
    Returns a list of font keywords that were not found.
    """
    if platform.system() != 'Windows':
        return []

    available_names = set()
    for font_dir in WINDOWS_FONT_DIRS:
        if not font_dir.exists():
            continue
        for pattern in ('*.ttf', '*.otf'):
            for font_path in font_dir.glob(pattern):
                available_names.add(font_path.stem.lower())
                available_names.add(font_path.name.lower())

    missing = []
    for keyword in REQUIRED_FONT_KEYWORDS:
        keyword_lower = keyword.lower()
        if not any(keyword_lower in name for name in available_names):
            missing.append(keyword)
    return missing

def setup_template():
    """
    HÆ°á»›ng dáº«n táº¡o template náº¿u chÆ°a cÃ³
    """
    if not os.path.exists(TEMPLATE_FILE):
        print("âš ï¸  CHÆ¯A CÃ“ TEMPLATE!")
        print("\nðŸ“ HÆ¯á»šNG DáºªN Táº O TEMPLATE TRONG WORD:")
        print("="*50)
        print("1. Má»Ÿ file Word cÃ³ background cá»§a báº¡n")
        print("2. Táº¡i vá»‹ trÃ­ cáº§n Ä‘iá»n thÃ´ng tin, gÃµ:")
        print("   {{prefix}}     <- Cho danh xÆ°ng (Mrs./Mr.)")
        print("   {{name}}       <- Cho tÃªn")  
        print("   {{position}}   <- Cho chá»©c vá»¥")
        print("\n3. Format text:")
        print("   - Select {{prefix}} â†’ Font Faustina, 12pt, mÃ u #EC2829")
        print("   - Select {{name}} â†’ Font Faustina Bold, 17pt, mÃ u #EC2829")
        print("   - Select {{position}} â†’ Font Faustina, 13pt, mÃ u #014F9D")
        print("\n4. Save as 'badge_template.docx'")
        print("="*50)
        return False
    return True

def replace_text_in_paragraph(paragraph, placeholders):
    """
    Replace placeholders in a paragraph while preserving formatting.
    Works by replacing text in runs without changing their style.
    """
    # Build full paragraph text
    full_text = ''.join(run.text for run in paragraph.runs)

    # Replace ALL placeholders in the text
    modified = False
    for placeholder, value in placeholders.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)
            modified = True

    # If any replacement was made, update the runs
    if modified:
        # Clear all runs and put new text in first run to preserve its formatting
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = full_text
            else:
                run.text = ''

def replace_in_textboxes(doc, placeholders):
    """
    Replace placeholders in textboxes while preserving formatting for each run.
    This approach replaces text within each run individually to keep formatting like bold.
    """
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    # Find all textbox content elements
    textboxes = doc.element.body.findall('.//w:txbxContent', namespaces)

    for textbox in textboxes:
        # Get all paragraphs within the textbox
        paras = textbox.findall('.//w:p', namespaces)
        for para_elem in paras:
            # Get all text runs
            runs = para_elem.findall('.//w:r', namespaces)

            # Process each run individually to preserve formatting
            for run in runs:
                t_elems = run.findall('.//w:t', namespaces)
                for t in t_elems:
                    if t.text:
                        # Replace placeholders in this specific text element
                        original = t.text
                        for placeholder, value in placeholders.items():
                            if placeholder in original:
                                t.text = original.replace(placeholder, value)
                                original = t.text

def create_badge_from_template(data, index):
    """
    Táº¡o badge tá»« template cho 1 ngÆ°á»i
    
    Args:
        data: Dictionary chá»©a thÃ´ng tin (prefix, name, position)
        index: Sá»‘ thá»© tá»±
    
    Returns:
        Path to PNG file
    """
    # Load template
    doc = Document(TEMPLATE_FILE)

    # Create placeholders mapping
    placeholders = {
        '{{prefix}}': data['prefix'],
        '{{name}}': data['name'],
        '{{position}}': data['position']
    }

    # Replace in all paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, placeholders)

    # Replace in tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, placeholders)

    # Replace in textboxes (most important for this template!)
    replace_in_textboxes(doc, placeholders)

    # Save Word file táº¡m
    temp_docx = os.path.join(TEMP_DIR, f'temp_{index:03d}.docx')
    doc.save(temp_docx)
    
    # Convert Word â†’ PDF
    temp_pdf = temp_docx.replace('.docx', '.pdf')
    try:
        convert(temp_docx, temp_pdf)
    except Exception as e:
        print(f"âš ï¸  Lá»—i convert PDF: {e}")
        print("   Äang thá»­ phÆ°Æ¡ng Ã¡n dá»± phÃ²ng...")
        # Fallback: dÃ¹ng COM object trá»±c tiáº¿p
        return convert_via_com(temp_docx, data['name'], index)
    
    # Convert PDF â†’ PNG
    try:
        poppler_kwargs = {}
        if POPPLER_PATH:
            poppler_kwargs['poppler_path'] = POPPLER_PATH
        elif not getattr(create_badge_from_template, "_poppler_warned", False):
            print("Poppler path not found in project; falling back to system PATH.")
            create_badge_from_template._poppler_warned = True

        images = convert_from_path(temp_pdf, dpi=300, fmt='png', **poppler_kwargs)
        
        # Save PNG
        output_file = os.path.join(OUTPUT_DIR, 
                                  f"{index:03d}_{data['name'].replace(' ', '_')}.png")
        images[0].save(output_file, 'PNG', quality=95)
        
        # Cleanup temp files
        os.remove(temp_docx)
        os.remove(temp_pdf)
        
        return output_file
        
    except Exception as e:
        print(f"âš ï¸  Lá»—i convert PNG: {e}")
        return None

def convert_via_com(docx_path, name, index):
    """
    Backup: Convert trá»±c tiáº¿p trong Word qua COM (Windows only)
    """
    try:
        import win32com.client
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Open document
        doc = word.Documents.Open(os.path.abspath(docx_path))
        
        # Export as PNG
        output_file = os.path.join(OUTPUT_DIR, f"{index:03d}_{name.replace(' ', '_')}.png")
        doc.SaveAs2(os.path.abspath(output_file), FileFormat=19)  # PNG format
        
        doc.Close()
        word.Quit()
        
        return output_file
        
    except:
        print("   âŒ KhÃ´ng thá»ƒ dÃ¹ng Word COM (cáº§n Windows + Word)")
        return None

def generate_all_badges():
    """
    Main function - Äá»c Excel vÃ  táº¡o táº¥t cáº£ badges
    """
    # Kiá»ƒm tra template
    if not setup_template():
        return
    
    # Äá»c Excel
    try:
        df = pd.read_excel(EXCEL_FILE)
        print(f"âœ“ Äá»c Ä‘Æ°á»£c {len(df)} ngÆ°á»i tá»« Excel")
    except FileNotFoundError:
        print(f"\nâŒ KhÃ´ng tÃ¬m tháº¥y file '{EXCEL_FILE}'")
        create_sample_excel()
        return
    except Exception as e:
        print(f"âŒ Lá»—i Ä‘á»c Excel: {e}")
        return
    
    # Chuáº©n hÃ³a tÃªn cá»™t (há»— trá»£ tiáº¿ng Viá»‡t)
    column_mapping = {
        'họ và tên': 'name',
        'họ tên': 'name', 
        'tên': 'name',
        'name': 'name',
        'chức vụ': 'position',
        'chức danh': 'position',
        'position': 'position',
        'title': 'position',
        'danh xưng': 'prefix',
        'prefix': 'prefix',
        'mr/mrs': 'prefix'
    }
    
    # Rename columns theo mapping
    df.columns = [column_mapping.get(col.lower().strip(), col) for col in df.columns]
    
    # Kiá»ƒm tra cá»™t báº¯t buá»™c
    if 'name' not in df.columns or 'position' not in df.columns:
        print(f"âŒ Thiáº¿u cá»™t báº¯t buá»™c!")
        print(f"   Cá»™t hiá»‡n cÃ³: {list(df.columns)}")
        print(f"   Cáº§n cÃ³: 'name' vÃ  'position' (hoáº·c tiáº¿ng Viá»‡t tÆ°Æ¡ng Ä‘Æ°Æ¡ng)")
        return
    
    # Xá»­ lÃ½ tá»«ng ngÆ°á»i
    success = 0
    failed = 0
    
    print(f"\nðŸš€ Báº¯t Ä‘áº§u táº¡o {len(df)} badges...")
    print("-" * 50)
    
    for idx, row in df.iterrows():
        try:
            # Chuáº©n bá»‹ data
            name = str(row['name']).strip()
            position = str(row['position']).strip() if pd.notna(row['position']) else ''
            
            # XÃ¡c Ä‘á»‹nh prefix
            if 'prefix' in row and pd.notna(row['prefix']):
                prefix = str(row['prefix']).strip()
            else:
                # Auto detect cho tÃªn tiáº¿ng Viá»‡t
                name_lower = name.lower()
                female_indicators = ['thá»‹', 'thu', 'hÆ°Æ¡ng', 'háº±ng', 'hÃ ', 'lan', 
                                    'mai', 'phÆ°Æ¡ng', 'thÃºy', 'thá»§y', 'trinh', 'trang']
                prefix = 'Mrs.' if any(ind in name_lower for ind in female_indicators) else 'Mr.'
            
            # Data cho template
            data = {
                'prefix': prefix,
                'name': name,
                'position': position
            }
            
            # Táº¡o badge
            output_file = create_badge_from_template(data, idx + 1)
            
            if output_file:
                success += 1
                print(f"  [{idx+1:3d}] âœ“ {name:<30} â†’ {os.path.basename(output_file)}")
            else:
                failed += 1
                print(f"  [{idx+1:3d}] âœ— {name:<30} â†’ Lá»—i táº¡o file")
                
        except Exception as e:
            failed += 1
            print(f"  [{idx+1:3d}] âœ— Lá»—i: {e}")
    
    # Cleanup temp folder
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)
    
    # BÃ¡o cÃ¡o káº¿t quáº£
    print("-" * 50)
    print(f"\nðŸ“Š Káº¾T QUáº¢:")
    print(f"  âœ“ ThÃ nh cÃ´ng: {success}/{len(df)}")
    if failed > 0:
        print(f"  âœ— Tháº¥t báº¡i: {failed}")
    print(f"  ðŸ“ Output: {os.path.abspath(OUTPUT_DIR)}/")
    print("\nâœ… HOÃ€N Táº¤T! Badges sáºµn sÃ ng Ä‘á»ƒ in.")

def create_sample_excel():
    """
    Táº¡o file Excel máº«u
    """
    sample_data = {
        'Há» vÃ  tÃªn': [
            'Phan Thá»‹ Nhi HÃ ',
            'Nguyá»…n VÄƒn An',
            'Tráº§n Thá»‹ Mai Lan', 
            'LÃª Äá»©c Minh',
            'HoÃ ng Thu Trang'
        ],
        'Chá»©c vá»¥': [
            'TrÆ°á»Ÿng Ban Kiá»ƒm SoÃ¡t',
            'GiÃ¡m Äá»‘c Äiá»u HÃ nh',
            'Káº¿ ToÃ¡n TrÆ°á»Ÿng',
            'TrÆ°á»Ÿng PhÃ²ng IT', 
            'PhÃ³ GiÃ¡m Äá»‘c Kinh Doanh'
        ],
        'Danh xÆ°ng': [
            'Mrs.', 'Mr.', 'Ms.', 'Mr.', 'Ms.'
        ]
    }
    
    df = pd.DataFrame(sample_data)
    df.to_excel(EXCEL_FILE, index=False)
    print(f"âœ“ ÄÃ£ táº¡o file Excel máº«u: {EXCEL_FILE}")
    print("  â†’ Cáº­p nháº­t file nÃ y vá»›i danh sÃ¡ch tháº­t")

def check_dependencies():
    """
    Kiá»ƒm tra vÃ  hÆ°á»›ng dáº«n cÃ i Ä‘áº·t dependencies
    """
    missing = []
    
    # Check Python packages
    try:
        import pandas
    except:
        missing.append('pandas')
    
    try:
        import docx
    except:
        missing.append('python-docx')
    
    try:
        import docx2pdf
    except:
        missing.append('docx2pdf')
        
    try:
        import pdf2image
    except:
        missing.append('pdf2image')
    
    if missing:
        print("âš ï¸  THIáº¾U THÆ¯ VIá»†N!")
        print(f"\nðŸ“¦ CÃ i Ä‘áº·t báº±ng lá»‡nh:")
        print(f"   pip install {' '.join(missing)} openpyxl")
        
        if 'pdf2image' in missing:
            print(f"\nðŸ“Œ LÆ°u Ã½: pdf2image cáº§n Poppler")
            print(f"   Windows: Táº£i tá»« https://github.com/oschwartz10612/poppler-windows/releases")
            print(f"   Mac: brew install poppler")
            print(f"   Linux: sudo apt-get install poppler-utils")
        
        return False
    
    # Check for Word (Windows)
    if platform.system() == 'Windows':
        missing_fonts = find_missing_fonts()
        if missing_fonts:
            print("WARNING: Faustina font files were not found on this machine.")
            print("         Install the Faustina family (Regular and Bold) so the badges match the reference design.")
            print("         Download: https://github.com/google/fonts/tree/main/ofl/faustina")
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Quit()
        except:
            print("âš ï¸  Microsoft Word khÃ´ng tÃ¬m tháº¥y hoáº·c pywin32 chÆ°a cÃ i")
            print("   pip install pywin32")
    
    return True

# ===== CHÆ¯Æ NG TRÃŒNH CHÃNH =====
if __name__ == '__main__':
    print("="*60)
    print("   WORD TEMPLATE BADGE GENERATOR - VIMC")
    print("   Font Faustina + MÃ u chuáº©n + Background Word")
    print("="*60)
    
    # Kiá»ƒm tra dependencies
    if not check_dependencies():
        print("\nâŒ Vui lÃ²ng cÃ i Ä‘áº·t cÃ¡c thÆ° viá»‡n cáº§n thiáº¿t trÆ°á»›c!")
        exit(1)
    
    # Cháº¡y
    try:
        generate_all_badges()
    except KeyboardInterrupt:
        print("\n\nâš ï¸  ÄÃ£ dá»«ng theo yÃªu cáº§u.")
    except Exception as e:
        print(f"\nâŒ Lá»—i khÃ´ng mong Ä‘á»£i: {e}")
        import traceback
        traceback.print_exc()

